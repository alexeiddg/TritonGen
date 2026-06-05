from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("poster_hallazgo_principal_tritongen.docx")
DOWNLOADS_DIR = Path("/Users/sebis/Downloads")
TEC_LOGO = Path("assets/tec_logo.png")
TODAY_ES = "Miércoles 3 de junio del 2026"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_check(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run("☑ ").bold = True
    p.add_run(text)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hallazgo principal para el póster")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("TritonGen - Evaluación experimental de generación de kernels con LLMs")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_cover_line(
    doc: Document,
    text: str,
    *,
    size: int = 14,
    bold: bool = False,
    before: int = 0,
    after: int = 8,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(120)

    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.paragraph_format.space_after = Pt(110)
    if TEC_LOGO.exists():
        logo.add_run().add_picture(str(TEC_LOGO), width=Inches(6.6))

    add_cover_line(doc, "Hallazgo principal para el póster", bold=True, after=28)
    add_cover_line(doc, "Desarrollo e implantación de sistemas de software", after=28)
    add_cover_line(doc, "Grupo 504", after=24)
    add_cover_line(doc, "Alexei Delgado De Gante | A01637405", after=2)
    add_cover_line(doc, "Luis Fernando Cuevas Arroyo | A01647254", after=2)
    add_cover_line(doc, "Aaron Hernandez Jimenez | A01642529", after=2)
    add_cover_line(doc, "Sebastian Cervera Maltos | A01068436", after=28)
    add_cover_line(doc, TODAY_ES, after=0)
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(text)


def add_graph(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_title(doc)

    doc.add_heading("Hallazgo principal", level=1)
    doc.add_paragraph(
        "El resultado más fuerte no fue que la gramática o el repair loop mejoraran "
        "claramente al modelo pequeño, sino que Claude compiló muchos más kernels "
        "Triton que Qwen y Gemini usando el mismo prompt, mientras que Gemini no "
        "logró ninguna compilación exitosa. Esto sugiere que en generación de kernels "
        "Triton no basta con usar un modelo grande: también importan el alineamiento "
        "del modelo con código Triton, el formato del prompt y la evaluación por etapas."
    )

    add_callout(
        doc,
        "Headline recomendado",
        "Más grande no siempre significa mejor: Claude compiló 87/180 kernels, "
        "Gemini 0/180 con el mismo prompt.",
        fill="E8EEF5",
    )

    doc.add_heading("Tabla de hallazgos", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(3.05), Inches(3.45)]
    set_table_width(table, widths)

    headers = ["Resultado observado", "Evidencia"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(11, 37, 69)
    set_repeat_table_header(table.rows[0])

    rows = [
        (
            "Claude superó ampliamente al modelo pequeño en compilación L1.",
            "Claude: 87/180 = 48.33% vs mejor condición Qwen G+C: 4/177 = 2.26%; "
            "diferencia aproximada de +46 puntos porcentuales; p < 0.001.",
        ),
        (
            "Gemini tuvo 0 compilaciones exitosas con el mismo prompt usado para Claude.",
            "Gemini: 0/180 = 0%; 88.9% de sus salidas fallaron en F0_PARSE.",
        ),
        (
            "Mejorar compilación no garantiza correctitud funcional.",
            "En condiciones SLM: functional_success = 0/714. Además, 4 casos de G+C "
            "compilaron pero fallaron en L2 con NaN.",
        ),
    ]
    for observed, evidence in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (observed, evidence), strict=True):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].add_run(text)
    set_table_width(table, widths)

    doc.add_heading("Respuesta a la hipótesis principal", level=1)
    doc.add_paragraph("No completamente.")
    doc.add_paragraph(
        "La hipótesis principal era sobre functional_success L2 en el ablation interno. "
        "Sin embargo, todas las condiciones SLM tuvieron functional_success = 0/714, "
        "por lo que no hubo variabilidad suficiente para probar la hipótesis principal."
    )
    doc.add_paragraph(
        "Sí responde una hipótesis secundaria: Claude supera significativamente al SLM "
        "en compilación L1."
    )

    doc.add_heading("Evidencia estadística o descriptiva disponible", level=1)
    for item in ["Media / porcentaje", "Visualización", "p-value", "IC95%", "Tamaño de efecto"]:
        add_check(doc, item)

    doc.add_heading("Gráficas recomendadas para el póster", level=1)
    graph_items = [
        "Tasa de compilación L1 por condición: none, G, G+C, Claude y Gemini. Debe ser la gráfica principal.",
        "compile@k por condición: compile@1, compile@5 y compile@10 para mostrar la probabilidad práctica de obtener al menos un kernel compilable.",
        "Distribución de failure codes por condición/modelo: explica si el problema fue formato, runtime o correctitud.",
        "Compile_success por kernel class: muestra que Claude domina en elementwise y reduction, pero matmul sigue siendo difícil para todos.",
    ]
    for item in graph_items:
        add_bullet(doc, item)

    doc.add_heading("Historia del póster", level=1)
    story = [
        (
            "Esperábamos que",
            "Las restricciones gramaticales y el feedback de correctitud ayudaran al modelo pequeño a generar kernels Triton funcionalmente correctos.",
        ),
        (
            "Pero encontramos que",
            "El modelo pequeño no produjo ningún kernel funcionalmente correcto en L2, y la mejora en compilación fue muy baja. En cambio, Claude logró 48.33% de compilación L1, mientras que Gemini obtuvo 0% usando el mismo prompt.",
        ),
        (
            "Esto es interesante porque",
            "Muestra que la capacidad del modelo y su alineamiento con generación de código Triton pesan más que los controles del pipeline en este tier de desarrollo. También muestra que dos modelos frontier pueden comportarse de forma radicalmente distinta con el mismo prompt.",
        ),
        (
            "Esto importa porque",
            "Para generar kernels GPU útiles, no basta con que el código parezca correcto ni con que compile. Se necesita evaluar por etapas: formato, compilación y correctitud numérica.",
        ),
    ]
    for label, text in story:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + ": ")
        r.bold = True
        p.add_run(text)

    add_callout(
        doc,
        "Historia final",
        "Claude fue el único modelo que logró una tasa alta de compilación, pero ningún "
        "resultado permite afirmar correctitud funcional todavía. La generación de kernels "
        "Triton requiere modelos fuertes, prompts alineados y evaluación por niveles.",
        fill="F4F6F9",
    )

    doc.add_page_break()
    doc.add_heading("Gráficas listas para el póster", level=1)
    doc.add_paragraph(
        "Estas visualizaciones resumen el hallazgo principal y pueden colocarse directamente "
        "en el póster. La primera gráfica debe ocupar el lugar principal; las demás funcionan "
        "como evidencia de apoyo."
    )
    add_graph(
        doc,
        DOWNLOADS_DIR / "WhatsApp Image 2026-06-03 at 21.55.01.jpeg",
        "Figura 1. Tasa de compilación L1 por condición con intervalos Wilson al 95%.",
    )
    add_graph(
        doc,
        DOWNLOADS_DIR / "WhatsApp Image 2026-06-03 at 21.55.21.jpeg",
        "Figura 2. Distribución de códigos de falla por condición/modelo.",
    )
    add_graph(
        doc,
        DOWNLOADS_DIR / "WhatsApp Image 2026-06-03 at 21.55.33.jpeg",
        "Figura 3. Compilación L1 por tipo de kernel.",
    )
    add_graph(
        doc,
        DOWNLOADS_DIR / "WhatsApp Image 2026-06-03 at 21.55.12.jpeg",
        "Figura 4. Probabilidad de obtener al menos un kernel compilable en k intentos.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
